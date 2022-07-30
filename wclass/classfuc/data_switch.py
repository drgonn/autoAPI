"""转换对象对象的方法"""
import os

from docx import Document
from docx.shared import Inches
from docx.shared import Cm
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

from wclass.const import ResetAPI


def make_table(document, table_list, merge_case=[]):
    """
    生成一张表格
    document: 文本
    table_list: 表结构二维数组,如[[1,2],[3,4]]
    merge_case: 需要合并的单元格
    """
    rows = len(table_list)
    columns = len(table_list[0])
    table = document.add_table(rows, columns, style="Table Grid")
    table.allow_autofit = True
    # table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for index_row, r in enumerate(table_list):
        for i, grid in enumerate(r):
            cells = table.rows[index_row].cells
            cells[i].text = grid
    for a, b, c, d in merge_case:
        table.rows[a].cells[b].merge(table.rows[c].cells[d])


def make_center_line(doc, title):
    p = doc.add_paragraph(title)
    paragraph_format = p.paragraph_format
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER


def make_center_title(doc, title, level):
    p = doc.add_heading(title, level=level)
    paragraph_format = p.paragraph_format
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER


def make_table_with_title(title, document, table_list, merge_case=[]):
    """
    生成一张表格带标题
    document: 文本
    table_list: 表结构二维数组,如[[1,2],[3,4]]
    merge_case: 需要合并的单元格
    """
    make_center_line(document, title)
    make_table(document, table_list, merge_case)


def to_doc(p):
    """将对象生成doc文档"""
    print("生成doc文档中...")
    print(p.doc_dir)
    doc = Document()
    # 设置字体
    doc.styles['Normal'].font.name = u'宋体'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    doc.styles['Normal'].font.size = Pt(10.5)
    doc.styles['Normal'].font.color.rgb = RGBColor(0, 0, 0)
    make_center_title(doc, f'{p.zh}功能设计说明书', 1)

    make_center_title(doc, '修订记录', 1)
    make_table(doc, [
        ["日期", "版本", "概述", "修订人", "审核人"],
        [p.now_date, "0.1.0", "初稿设计", "谌榕", ""],
        ["", "", "", "", ""],
    ])

    doc.add_heading("一、功能说明", level=1)
    doc.add_heading("二、数据库设计", level=1)
    # doc.add_heading("2.1 数据库设计", level=3)
    doc.add_paragraph("mysql数据库设计说明")
    table_list = [
        ["表名", "字段", "类型", "说明"],
    ]
    for t in p.tables:
        table_list.append([t.names, "", "", t.zh_name+t.about])
        for c in t.columns:
            table_list.append(["", c.name, c.sql_type, c.zh])
    make_table(doc, table_list)

    doc.add_heading("三、HTTP接口说明", level=1)
    for i, t in enumerate(p.tables):
        doc.add_heading(f"3.{i+1} {t.zh_name}相关接口", level=3)
        for j, api in enumerate(ResetAPI):
            doc.add_paragraph(f"3.{i+1}.{j+1} {api.zh}{t.zh_name}接口")
            index = "/{" + t.index_name + "}" if api.index else ""
            doc.add_paragraph(f"    接口地址：{t.url_prefix}/{t.names}{index}")
            doc.add_paragraph(f"    YAPI测试地址：")
            doc.add_paragraph(f"    请求方式：{api.METHOD}")
            table_list = [
                ["参数", "含义", "类型", "必须", "说明"],
            ]
            if api.zh == "创建":
                for c in t.columns:
                    if c.post:
                        must = "是" if c.post == 2 else "否"
                        table_list.append([c.name, c.zh, c.db, must, c.about])
                make_table(doc, table_list)
            elif api.zh == '修改':
                for c in t.columns:
                    if c.put:
                        must = "是" if c.put == 2 else "否"
                        table_list.append([c.name, c.zh, c.db, must, c.about])
                make_table(doc, table_list)
            elif api.zh == '单个查询':
                pass
            elif api.zh == '列表查询':
                for c in t.columns:
                    if c.list:
                        find = "模糊查找参数" if c.put == 2 else "精确查找"
                        table_list.append([c.name, c.zh, c.db, "否", find])
                make_table(doc, table_list)
            elif api.zh == '单个删除':
                pass
            elif api.zh == '批量删除':
                table_list.append(
                    [t.index_name + "s", t.index_name + "数组", "Array", "是", "批量删除用的唯一标识列表"])
                make_table(doc, table_list)

    doc.add_heading("四、影响分析", level=1)
    doc.add_heading("五、计划", level=1)
    doc.add_heading("六、质量目标", level=1)
    doc.add_heading("七、测试建议", level=1)
    doc.add_heading("八、其他信息", level=1)
    doc.add_heading("九、参考资料", level=1)

    doc_dir = os.path.join(p.doc_dir, f"{p.zh}功能设计说明书-V0.1.0.docx")
    doc.save(doc_dir)
    print("生成文档成功", doc_dir)
    return doc_dir


def doc2class(p):
    """将doc文档生成对象"""
    #打开word
    doc_dir = os.path.join(p.doc_dir, f"{p.zh}功能设计说明书-V0.1.0.docx")
    doc = Document(doc_dir)
    for i in range(len(doc.paragraphs)):
        print(doc.paragraphs[i].text)